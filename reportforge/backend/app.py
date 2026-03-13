from __future__ import annotations
import json, uuid, datetime as dt
from pathlib import Path
from typing import Any, Optional

from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from sqlalchemy import select

from .db import engine, SessionLocal, Base
from .models import Report, Template, Finding

app = FastAPI(
    title='ReportForge API',
    version='1.2.5')

app.add_middleware(
    CORSMiddleware,
    allow_origins=['*'],
    allow_credentials=False,
    allow_methods=['*'],
    allow_headers=['*'],
)

_FRONTEND_DIR = (
    Path(__file__).resolve()
    .parents[1] / 'frontend'
).resolve()

def _migrate_findings_cols():
    cols = [
        ('discussion', 'TEXT'),
        ('refs', 'TEXT'),
    ]
    from sqlalchemy import text
    with engine.connect() as con:
        for col, typ in cols:
            sql = (
                'ALTER TABLE findings'
                ' ADD COLUMN '
                + col + ' ' + typ
                + ' DEFAULT ""'
            )
            try:
                con.execute(text(sql))
                con.commit()
            except Exception:
                pass

@app.on_event('startup')
def _startup():
    from reportforge.utils import get_db_path
    import logging
    logging.basicConfig(level=logging.INFO)
    logging.getLogger('reportforge').info(
        'DB: %s', get_db_path()
    )
    Base.metadata.create_all(bind=engine)
    _migrate_findings_cols()

@app.get('/', response_class=HTMLResponse)
def serve_frontend():
    return (
        _FRONTEND_DIR / 'index.html'
    ).read_text(encoding='utf-8')

def _ts(d: dt.datetime) -> str:
    return (
        d.replace(microsecond=0)
        .isoformat() + 'Z'
    )

def _now_iso() -> str:
    return (
        dt.datetime.utcnow()
        .replace(microsecond=0)
        .isoformat() + 'Z'
    )

@app.get('/api/health')
def health():
    from reportforge.utils import get_db_path
    return {
        'ok': True,
        'time': _now_iso(),
        'db': str(get_db_path()),
    }

class ReportCreate(BaseModel):
    title: str = 'Untitled Report'
    org: str = ''
    report_type: str = Field(
        default='Security Assessment',
        alias='type')
    classification: str = 'Confidential'
    assessment_date: str = Field(
        default='', alias='date')
    authors: str = ''
    template_id: Optional[str] = None
    model_config = {
        'populate_by_name': True}

class ReportUpdate(BaseModel):
    title: Optional[str] = None
    org: Optional[str] = None
    report_type: Optional[str] = Field(
        default=None, alias='type')
    classification: Optional[str] = None
    assessment_date: Optional[str] = Field(
        default=None, alias='date')
    authors: Optional[str] = None
    status: Optional[str] = None
    data: Optional[dict[str, Any]] = None
    model_config = {
        'populate_by_name': True}

class FindingCreate(BaseModel):
    report_id: Optional[str] = None
    title: str = ''
    severity: str = 'Medium'
    status: str = 'open'
    description: str = ''
    discussion: str = ''
    recommendation: str = ''
    refs: str = ''
    cvss: Optional[float] = None

class FindingUpdate(BaseModel):
    report_id: Optional[str] = None
    title: Optional[str] = None
    severity: Optional[str] = None
    status: Optional[str] = None
    description: Optional[str] = None
    discussion: Optional[str] = None
    recommendation: Optional[str] = None
    refs: Optional[str] = None
    cvss: Optional[float] = None

class TemplateCreate(BaseModel):
    name: str
    description: str = ''
    sections: list[dict[str, Any]] = []
    framework_id: Optional[str] = None

class TemplateUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    sections: Optional[
        list[dict[str, Any]]] = None

@app.get('/api/reports')
def list_reports(
    status: Optional[str] = None):
    with SessionLocal() as db:
        stmt = select(Report).order_by(
            Report.updated_at.desc())
        if status:
            stmt = stmt.where(
                Report.status == status)
        rows = (
            db.execute(stmt)
            .scalars().all())
        return [_report_out(r) for r in rows]

@app.post('/api/reports', status_code=201)
def create_report(payload: ReportCreate):
    rid = str(uuid.uuid4())
    init: dict[str, Any] = {
        'report': {
            'title': payload.title,
            'org': payload.org,
            'type': payload.report_type,
            'classification':
                payload.classification,
            'date': payload.assessment_date,
            'authors': payload.authors,
            'summary': '', 'scope': '',
        },
        'sections': [], 'findings': [],
    }
    if payload.template_id:
        with SessionLocal() as db:
            t = db.get(
                Template, payload.template_id)
            if t:
                tmpl = json.loads(
                    t.template_json or '{}')
                init['sections'] = (
                    tmpl.get('sections', []))
                init['template_id'] = (
                    payload.template_id)
                init['template_name'] = t.name
    with SessionLocal() as db:
        r = Report(
            id=rid,
            title=(
                payload.title.strip()
                or 'Untitled Report'),
            org=payload.org,
            report_type=payload.report_type,
            classification=(
                payload.classification),
            assessment_date=(
                payload.assessment_date),
            authors=payload.authors,
            status='draft',
            data_json=json.dumps(init),
        )
        db.add(r); db.commit()
    return {'id': rid}

@app.get('/api/reports/{rid}')
def get_report(rid: str):
    with SessionLocal() as db:
        stmt = select(Report).where(
            Report.id == rid)
        r = (
            db.execute(stmt)
            .scalar_one_or_none())
        if not r:
            raise HTTPException(
                404,
                f'Report not found: {rid!r}')
        return _report_out(r, include_data=True)

@app.put('/api/reports/{rid}')
def update_report(
    rid: str, payload: ReportUpdate):
    with SessionLocal() as db:
        stmt = select(Report).where(
            Report.id == rid)
        r = (
            db.execute(stmt)
            .scalar_one_or_none())
        if not r:
            raise HTTPException(
                404,
                f'Report not found: {rid!r}')
        if payload.title is not None:
            r.title = (
                payload.title.strip() or r.title)
        if payload.org is not None:
            r.org = payload.org
        if payload.report_type is not None:
            r.report_type = payload.report_type
        if payload.classification is not None:
            r.classification = (
                payload.classification)
        if payload.assessment_date is not None:
            r.assessment_date = (
                payload.assessment_date)
        if payload.authors is not None:
            r.authors = payload.authors
        if payload.status is not None:
            r.status = payload.status
        if payload.data is not None:
            rep = payload.data.get(
                'report', {})
            if rep.get('title'):
                r.title = rep['title']
            if rep.get('org') is not None:
                r.org = rep['org']
            r.data_json = json.dumps(
                payload.data)
        r.updated_at = dt.datetime.utcnow()
        db.add(r)
        db.commit()
        db.refresh(r)
        return {
            'ok': True,
            'updated_at': _ts(r.updated_at)
        }

@app.delete('/api/reports/{rid}')
def delete_report(rid: str):
    with SessionLocal() as db:
        r = db.get(Report, rid)
        if not r:
            raise HTTPException(
                404, 'Report not found')
        db.delete(r)
        db.commit()
    return {'ok': True}

def _report_out(
    r: Report,
    include_data: bool = False) -> dict:
    out = {
        'id': r.id,
        'title': r.title,
        'org': r.org,
        'type': r.report_type,
        'classification': r.classification,
        'date': r.assessment_date,
        'authors': r.authors,
        'status': r.status,
        'created_at': _ts(r.created_at),
        'updated_at': _ts(r.updated_at),
    }
    if include_data:
        try:
            out['data'] = json.loads(
                r.data_json or '{}')
        except:
            out['data'] = {}
    return out

@app.get('/api/findings')
def list_findings(
    report_id: Optional[str] = None):
    with SessionLocal() as db:
        stmt = select(Finding).order_by(
            Finding.updated_at.desc())
        if report_id:
            stmt = stmt.where(
                Finding.report_id == report_id)
        rows = (
            db.execute(stmt)
            .scalars().all())
        return [_finding_out(f) for f in rows]

@app.post('/api/findings', status_code=201)
def create_finding(
    payload: FindingCreate):
    fid = str(uuid.uuid4())
    with SessionLocal() as db:
        f = Finding(
            id=fid,
            report_id=payload.report_id,
            title=payload.title,
            severity=payload.severity,
            status=payload.status,
            description=payload.description,
            discussion=payload.discussion,
            recommendation=(
                payload.recommendation),
            refs=payload.refs,
            cvss=payload.cvss,
        )
        db.add(f)
        db.commit()
        db.refresh(f)
        return _finding_out(f)

@app.get('/api/findings/{fid}')
def get_finding(fid: str):
    with SessionLocal() as db:
        f = db.get(Finding, fid)
        if not f:
            raise HTTPException(
                404, 'Finding not found')
        return _finding_out(f)

@app.put('/api/findings/{fid}')
def update_finding(
    fid: str, payload: FindingUpdate):
    with SessionLocal() as db:
        f = db.get(Finding, fid)
        if not f:
            raise HTTPException(
                404, 'Finding not found')
        if payload.report_id is not None:
            f.report_id = payload.report_id
        if payload.title is not None:
            f.title = payload.title
        if payload.severity is not None:
            f.severity = payload.severity
        if payload.status is not None:
            f.status = payload.status
        if payload.description is not None:
            f.description = payload.description
        if payload.discussion is not None:
            f.discussion = payload.discussion
        if payload.recommendation is not None:
            f.recommendation = (
                payload.recommendation)
        if payload.refs is not None:
            f.refs = payload.refs
        if payload.cvss is not None:
            f.cvss = payload.cvss
        f.updated_at = dt.datetime.utcnow()
        db.add(f)
        db.commit()
        db.refresh(f)
        return _finding_out(f)

@app.delete('/api/findings/{fid}')
def delete_finding(fid: str):
    with SessionLocal() as db:
        f = db.get(Finding, fid)
        if not f:
            raise HTTPException(
                404, 'Finding not found')
        db.delete(f)
        db.commit()
    return {'ok': True}

def _finding_out(f: Finding) -> dict:
    return {
        'id': f.id,
        'report_id': f.report_id,
        'title': f.title,
        'severity': f.severity,
        'status': f.status,
        'description': f.description,
        'discussion': f.discussion,
        'recommendation': f.recommendation,
        'refs': f.refs,
        'cvss': f.cvss,
        'created_at': _ts(f.created_at),
        'updated_at': _ts(f.updated_at),
    }

@app.get('/api/templates')
def list_templates():
    with SessionLocal() as db:
        rows = (
            db.execute(
                select(Template).order_by(
                    Template.updated_at.desc()))
            .scalars().all())
        return [_template_out(t) for t in rows]

@app.post('/api/templates', status_code=201)
def create_template(payload: TemplateCreate):
    tid = str(uuid.uuid4())
    with SessionLocal() as db:
        t = Template(
            id=tid,
            name=(
                payload.name.strip()
                or 'Untitled Template'),
            description=payload.description,
            template_json=json.dumps(
                {'sections': payload.sections,
                 'framework_id': payload.framework_id}),
        )
        db.add(t)
        db.commit()
    return {'id': tid}

@app.get('/api/templates/{tid}')
def get_template(tid: str):
    with SessionLocal() as db:
        t = db.get(Template, tid)
        if not t:
            raise HTTPException(
                404, 'Template not found')
        return _template_out(t, include_data=True)

@app.put('/api/templates/{tid}')
def update_template(
    tid: str, payload: TemplateUpdate):
    with SessionLocal() as db:
        t = db.get(Template, tid)
        if not t:
            raise HTTPException(
                404, 'Template not found')
        if payload.name is not None:
            t.name = (
                payload.name.strip() or t.name)
        if payload.description is not None:
            t.description = payload.description
        if payload.sections is not None:
            existing = json.loads(
                t.template_json or '{}')
            existing['sections'] = (
                payload.sections)
            t.template_json = json.dumps(
                existing)
        t.updated_at = dt.datetime.utcnow()
        db.add(t)
        db.commit()
        db.refresh(t)
        return {
            'ok': True,
            'updated_at': _ts(t.updated_at)
        }

@app.delete('/api/templates/{tid}')
def delete_template(tid: str):
    with SessionLocal() as db:
        t = db.get(Template, tid)
        if not t:
            raise HTTPException(
                404, 'Template not found')
        db.delete(t)
        db.commit()
    return {'ok': True}

def _template_out(
    t: Template,
    include_data: bool = False) -> dict:
    try:
        _tj=json.loads(t.template_json or '{}')
    except:
        _tj={}
    out = {
        'id': t.id,
        'name': t.name,
        'description': t.description,
        'framework_id': _tj.get('framework_id'),
        'created_at': _ts(t.created_at),
        'updated_at': _ts(t.updated_at),
    }
    if include_data:
        try:
            d = json.loads(
                t.template_json or '{}')
            out['sections'] = d.get(
                'sections', [])
        except:
            out['sections'] = []
    return out

def _build_report_html(r, findings, branding_logo=''):
    t = r.title or 'Report'
    esc = lambda s: (
        str(s)
        .replace('&','&amp;')
        .replace('<','&lt;')
        .replace('>','&gt;')
    )
    d = json.loads(r.data_json or '{}')
    rep = d.get('report', {})
    logo_b64 = rep.get('logo','')
    org = esc(r.org or rep.get('org',''))
    assessor = esc(r.authors or rep.get('assessor',''))
    end_date = esc(rep.get('end_date','') or r.assessment_date or '')
    delivery_date = esc(rep.get('delivery_date',''))
    poc_first = rep.get('poc_first','')
    poc_last = rep.get('poc_last','')
    poc_email = rep.get('poc_email','')
    poc_phone = rep.get('poc_phone','')
    has_poc = any([poc_first, poc_last, poc_email, poc_phone])
    # --- build logo img tags ---
    def _img_tag(b64, max_h, max_w, alt):
        if not b64:
            return ''
        src = b64 if ',' in b64 else 'data:image/png;base64,' + b64
        return ('<img src="' + src + '" style="max-height:' + max_h
                + ';max-width:' + max_w + ';object-fit:contain" alt="' + alt + '">')
    branding_html = _img_tag(branding_logo, '80px', '240px', 'Logo')
    company_html = _img_tag(logo_b64, '80px', '200px', 'Company Logo')
    # --- title page ---
    submitted_block = ''
    if org or company_html:
        submitted_block = '<div class="tp-submitted-block">'
        if org:
            submitted_block += '<p class="tp-submitted-org">Submitted to: ' + org + '</p>'
        elif company_html:
            submitted_block += '<p class="tp-submitted-org">Submitted to:</p>'
        if company_html:
            submitted_block += '<div class="tp-company-logo">' + company_html + '</div>'
        submitted_block += '</div>'
    title_page = (
        '<div class="title-page">'
        '<div class="tp-top">'
        + (branding_html if branding_html else '')
        + '</div>'
        '<div class="tp-mid">'
        '<p class="tp-report-title">Security Assessment &#8211; Final Report</p>'
        + submitted_block
        + '</div>'
        + ('<div class="tp-bot">'
        + ('<p class="tp-meta-line">Prepared By: ' + assessor + '</p>' if assessor else '')
        + ('<p class="tp-meta-line">Date Issued: ' + delivery_date + '</p>' if delivery_date else '')
        + '</div>' if (assessor or delivery_date) else '')
        + '</div>'
    )
    # Build title-page footer (confidentiality notice only)
    title_footer_html = (
        '<div class="page-footer">'
        '<div class="footer-title-page">'
        'This document contains information that is confidential and privileged. '
        'Unless you are the intended recipient, you may not use, copy or disclose '
        'to anyone the information contained herein.'
        '</div>'
        '</div>'
    )
    # Build standard footer (Confidential | Page # | Company Name)
    std_footer_html = (
        '<div class="page-footer footer-std">'
        '<span>Confidential</span>'
        '<span id="pg-num"></span>'
        '<span>' + esc(org) + '</span>'
        '</div>'
        '<script>'
        '(function(){'
        'var pages=document.querySelectorAll(".report-body");'
        'var f=document.getElementById("pg-num");'
        'if(f)f.textContent="Page "+window.location.hash.replace("#","");'
        '})();'
        '</script>'
    )
    parts = [
        '<!DOCTYPE html>',
        '<html><head>',
        '<meta charset="UTF-8">',
        '<title>' + esc(t) + '</title>',
        '<style>',
        'body{font-family:Arial,sans-serif;',
        'margin:0;color:#1a1d27}',
        '.title-page{display:flex;flex-direction:column;',
        'justify-content:space-between;',
        'min-height:100vh;padding:60px 60px 60px 60px;',
        'box-sizing:border-box;page-break-after:always}',
        '.tp-top{text-align:center;padding-top:40px}',
        '.tp-mid{text-align:center;padding:40px 0}',
        '.tp-report-title{font-size:26pt;font-weight:700;',
        'color:#1a1d27;margin:0 0 20px 0;line-height:1.2}',
        '.tp-submitted-block{margin-top:20px}',
        '.tp-submitted-label{font-size:11pt;color:#666;text-transform:uppercase;letter-spacing:.05em;margin:0 0 8px 0}',
        '.tp-company-logo{margin:8px 0}',
        '.tp-submitted-org{font-size:14pt;color:#444;margin:8px 0 0 0}',
        '.tp-bot{margin-top:auto;padding-bottom:80px}',
        '.tp-meta-line{font-size:12pt;color:#333;margin:6px 0}',
        '.report-body{padding:40px 60px}',
        'h1{font-size:28px;margin-bottom:8px;margin-top:32px}',
        'h2{font-size:22px;margin-top:28px;',
        'border-bottom:2px solid #4f6ef7;',
        'padding-bottom:4px}',
        'h3{font-size:18px;margin-top:24px}',
        'h4{font-size:16px;margin-top:20px}',
        'h5{font-size:14px;margin-top:16px}',
        'h6{font-size:13px;margin-top:14px;color:#555}',
        '.finding{border:1px solid #ccc;',
        'border-radius:6px;padding:16px;',
        'margin-bottom:16px}',
        '.sev{display:inline-block;',
        'padding:2px 8px;border-radius:4px;',
        'font-size:11px;font-weight:700;',
        'margin-left:8px}',
        '.critical{background:#d32f2f;color:#fff}',
        '.high{background:#f57c00;color:#fff}',
        '.medium{background:#f9a825;color:#000}',
        '.low{background:#388e3c;color:#fff}',
        '.info{background:#0288d1;color:#fff}',
        '.sec-content{margin:8px 0 16px 0;',
        'line-height:1.6}',
        '@page{margin-bottom:60px}',
        '@page:first{margin-bottom:80px}',
        '.page-footer{position:fixed;bottom:0;left:0;right:0;',
        'padding:10px 60px;font-size:9pt;color:#555;',
        'border-top:1px solid #ccc;background:#fff}',
        '.footer-title-page{text-align:center;font-style:italic;',
        'font-size:8.5pt;color:#666;line-height:1.4}',
        '.footer-std{display:flex;justify-content:space-between;',
        'align-items:center}',
        '.footer-meta{margin-top:6px;font-size:9pt;color:#444}',
        '</style></head><body>',
        title_page,
        title_footer_html,
        std_footer_html,
        '<div class="report-body">',
    ]
    secs = d.get('sections', [])
    for sec in secs:
        hl = sec.get('heading_level', 1)
        hl = hl if isinstance(hl, int) and 1 <= hl <= 6 else 1
        htag = 'h' + str(hl)
        parts.append(
            '<' + htag + '>' + esc(sec.get(
            'title','')) + '</' + htag + '>')
        stitle = sec.get('title','').lower()
        if has_poc and ('appendix b' in stitle or 'points of contact' in stitle):
            poc_name = (esc(poc_first) + ' ' + esc(poc_last)).strip()
            parts.append(
                '<p style="margin:16px 0 4px 0">The primary point of contact for this assessment was:</p>'
                '<p style="margin:4px 0 2px 0"><strong>' + poc_name + '</strong></p>'
                + ('<p style="margin:2px 0">' + esc(poc_email) + '</p>' if poc_email else '')
                + ('<p style="margin:2px 0 16px 0">' + esc(poc_phone) + '</p>' if poc_phone else '')
            )
        content = sec.get('content','')
        if content:
            parts.append(
                '<div class="sec-content">'
                + content + '</div>')
    if findings:
        parts.append(
            '<h2>Findings (' +
            str(len(findings)) + ')</h2>')
        for f in findings:
            sev = (f.get('severity')
                   or 'info').lower()
            parts.append(
                '<div class="finding">')
            parts.append(
                '<b>' + esc(f.get(
                'title','')) + '</b>'
                + '<span class="sev '
                + sev + '">' + sev
                + '</span>')
            for lbl, key in [
                ('Observation','description'),
                ('Discussion','discussion'),
                ('Recommendations',
                 'recommendation'),
                ('References','refs'),
            ]:
                val = f.get(key,'')
                if val:
                    parts.append(
                        '<p><b>' + lbl
                        + ':</b></p>'
                        + '<div>' + val
                        + '</div>')
            parts.append('</div>')
    parts.append('</div>')  # close .report-body
    parts.append('</body></html>')
    return '\n'.join(parts)

def _get_report_for_export(rid):
    with SessionLocal() as db:
        r = db.get(Report, rid)
        if not r:
            raise HTTPException(
                404, 'Report not found')
        d = json.loads(
            r.data_json or '{}')
        findings = d.get('findings', [])
        return r, findings

class ExportBody(BaseModel):
    branding_logo: str = ''

@app.post('/api/reports/{rid}/export/html')
def export_html(rid: str, body: ExportBody = ExportBody()):
    r, findings = _get_report_for_export(rid)
    html = _build_report_html(r, findings, branding_logo=body.branding_logo)
    fname = (r.title or 'report').replace(' ', '_') + '.html'
    return Response(content=html, media_type='text/html',
        headers={'Content-Disposition': 'attachment; filename="' + fname + '"'})

@app.post('/api/reports/{rid}/export/pdf')
def export_pdf(rid: str, body: ExportBody = ExportBody()):
    r, findings = _get_report_for_export(rid)
    html = _build_report_html(r, findings, branding_logo=body.branding_logo)
    try:
        from weasyprint import HTML
        pdf = HTML(string=html).write_pdf()
    except Exception as e:
        raise HTTPException(500, 'PDF generation failed: ' + str(e))
    fname = (r.title or 'report').replace(' ', '_') + '.pdf'
    return Response(content=pdf, media_type='application/pdf',
        headers={'Content-Disposition': 'attachment; filename="' + fname + '"'})

def _add_html_to_docx(doc, html, base_heading=2):
    import re, base64, io as _io
    from bs4 import BeautifulSoup, NavigableString, Tag
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    BLOCK_TAGS = {'p','div','h1','h2','h3','h4',
                  'h5','h6','ul','ol','li','br',
                  'blockquote','pre'}
    H_MAP = {'h1': base_heading,
             'h2': base_heading,
             'h3': base_heading+1,
             'h4': base_heading+2,
             'h5': base_heading+2,
             'h6': base_heading+3}
    def parse_color(style_str):
        m = re.search(r'colors*:s*#([0-9a-fA-F]{6})', style_str or '')
        if m:
            h = m.group(1)
            return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))
        m2 = re.search(r'colors*:s*rgb((d+),s*(d+),s*(d+))', style_str or '')
        if m2:
            return RGBColor(int(m2.group(1)),int(m2.group(2)),int(m2.group(3)))
        return None
    def add_run_from_node(para, node, bold=False, italic=False, color=None):
        if isinstance(node, NavigableString):
            txt = str(node)
            if not txt:
                return
            run = para.add_run(txt)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color
            return
        if not isinstance(node, Tag):
            return
        tag = node.name.lower() if node.name else ''
        new_bold = bold or tag in ('b','strong')
        new_italic = italic or tag in ('i','em')
        new_color = color
        style_str = node.get('style','')
        c = parse_color(style_str)
        if c:
            new_color = c
        if tag == 'img':
            src = node.get('src','')
            if src.startswith('data:image'):
                try:
                    header, b64data = src.split(',',1)
                    img_bytes = base64.b64decode(b64data)
                    img_stream = _io.BytesIO(img_bytes)
                    from docx.shared import Inches
                    run = para.add_run()
                    run.add_picture(img_stream, width=Inches(4))
                except Exception:
                    pass
            return
        for child in node.children:
            add_run_from_node(para, child, new_bold, new_italic, new_color)
    def process_block(node):
        if isinstance(node, NavigableString):
            txt = str(node).strip()
            if txt:
                doc.add_paragraph(txt)
            return
        if not isinstance(node, Tag):
            return
        tag = node.name.lower() if node.name else ''
        if tag in H_MAP:
            doc.add_heading(node.get_text(), H_MAP[tag])
            return
        if tag == 'img':
            src = node.get('src','')
            if src.startswith('data:image'):
                try:
                    header, b64data = src.split(',',1)
                    img_bytes = base64.b64decode(b64data)
                    img_stream = _io.BytesIO(img_bytes)
                    from docx.shared import Inches
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_stream, width=Inches(4))
                except Exception:
                    pass
            return
        if tag in ('ul','ol'):
            list_style = 'List Bullet' if tag=='ul' else 'List Number'
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style=list_style)
                for child in li.children:
                    add_run_from_node(p, child)
            return
        if tag in ('p','div','blockquote','pre'):
            has_block = any(
                isinstance(c, Tag) and c.name and
                c.name.lower() in BLOCK_TAGS
                for c in node.children)
            if has_block:
                for child in node.children:
                    process_block(child)
            else:
                p = doc.add_paragraph()
                for child in node.children:
                    add_run_from_node(p, child)
            return
        if tag == 'br':
            doc.add_paragraph()
            return
        for child in node.children:
            process_block(child)
    soup = BeautifulSoup(html, 'html.parser')
    for child in soup.children:
        process_block(child)


@app.post('/api/reports/{rid}/export/docx')
def export_docx(rid: str, body: ExportBody = ExportBody()):
    r, findings = _get_report_for_export(rid)
    try:
        import io, base64, io as _bio
        from docx import Document
        from bs4 import BeautifulSoup
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        doc = Document()
        d = json.loads(r.data_json or '{}')
        rep = d.get('report', {})
        logo_b64 = rep.get('logo', '')
        branding_logo = body.branding_logo or ''
        org_name = r.org or rep.get('org', '')
        assessor_name = r.authors or rep.get('assessor', '')
        delivery_date_val = rep.get('delivery_date', '')
        poc_first = rep.get('poc_first', '')
        poc_last = rep.get('poc_last', '')
        poc_email = rep.get('poc_email', '')
        poc_phone = rep.get('poc_phone', '')
        has_poc = any([poc_first, poc_last, poc_email, poc_phone])

        def _add_img(b64, width_in):
            if not b64:
                return
            try:
                raw = b64.split(',', 1)[1] if ',' in b64 else b64
                stream = _bio.BytesIO(base64.b64decode(raw))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(stream, width=Inches(width_in))
            except Exception:
                pass

        # --- Title page ---
        # Top: site branding logo
        _add_img(branding_logo, 2.5)

        # Report title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run('Security Assessment – Final Report')
        run_title.bold = True
        run_title.font.size = Pt(26)

        # Submitted to: + company logo from report
        if org_name or logo_b64:
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub_label = p_sub.add_run('Submitted to:')
            run_sub_label.font.size = Pt(14)
            if org_name:
                run_sub_org = p_sub.add_run(' ' + org_name)
                run_sub_org.font.size = Pt(14)
            _add_img(logo_b64, 2.0)

        # Prepared By / Date Issued - stacked at bottom of title page, above footer
        if assessor_name or delivery_date_val:
            p_meta_anchor = doc.add_paragraph()
            p_meta_anchor.paragraph_format.space_before = Pt(360)
            p_meta_anchor.paragraph_format.space_after = Pt(0)
            if assessor_name:
                p_prep = doc.add_paragraph()
                p_prep.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_prep = p_prep.add_run('Prepared By: ' + assessor_name)
                run_prep.font.size = Pt(12)
            if delivery_date_val:
                p_date = doc.add_paragraph()
                p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_date = p_date.add_run('Date Issued: ' + delivery_date_val)
                run_date.font.size = Pt(12)

                # ── Footers ────────────────────────────────────────────────────────────
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        # First-page footer: confidentiality notice + Prepared By / Date Issued
        fp_footer = section.first_page_footer
        fp_footer.is_linked_to_previous = False
        p_conf = fp_footer.paragraphs[0] if fp_footer.paragraphs else fp_footer.add_paragraph()
        p_conf.clear()
        p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_conf = p_conf.add_run(
            'This document contains information that is confidential and privileged. '
            'Unless you are the intended recipient, you may not use, copy or disclose '
            'to anyone the information contained herein.')
        run_conf.font.size = Pt(8)
        run_conf.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_conf.font.italic = True


        # Standard footer: Confidential | Page # | Company Name
        std_footer = section.footer
        std_footer.is_linked_to_previous = False
        p_std = std_footer.paragraphs[0] if std_footer.paragraphs else std_footer.add_paragraph()
        p_std.clear()
        from docx.oxml import OxmlElement as _OE
        from docx.oxml.ns import qn as _qn
        tbl = _OE('w:tbl')
        tbl_pr = _OE('w:tblPr')
        tbl_w = _OE('w:tblW')
        tbl_w.set(_qn('w:w'), '9360')
        tbl_w.set(_qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)
        tbl_borders = _OE('w:tblBorders')
        for border_name in ['top','left','bottom','right','insideH','insideV']:
            b = _OE('w:' + border_name)
            b.set(_qn('w:val'), 'none')
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)
        tbl.append(tbl_pr)
        row = _OE('w:tr')
        def _footer_cell(text_or_fld, align, width, use_page_field=False):
            tc = _OE('w:tc')
            tcp = _OE('w:tcPr')
            tcw = _OE('w:tcW')
            tcw.set(_qn('w:w'), str(width))
            tcw.set(_qn('w:type'), 'dxa')
            tcp.append(tcw)
            tc.append(tcp)
            p = _OE('w:p')
            pp = _OE('w:pPr')
            jc = _OE('w:jc')
            jc.set(_qn('w:val'), align)
            pp.append(jc)
            p.append(pp)
            rpr = _OE('w:rPr')
            sz = _OE('w:sz')
            sz.set(_qn('w:val'), '18')
            rpr.append(sz)
            if use_page_field:
                rt = _OE('w:r')
                rt.append(rpr)
                t = _OE('w:t')
                t.set(_qn('xml:space'), 'preserve')
                t.text = 'Page '
                rt.append(t)
                p.append(rt)
                r2 = _OE('w:r')
                fld_begin = _OE('w:fldChar')
                fld_begin.set(_qn('w:fldCharType'), 'begin')
                r2.append(fld_begin)
                p.append(r2)
                r3 = _OE('w:r')
                instr = _OE('w:instrText')
                instr.set(_qn('xml:space'), 'preserve')
                instr.text = ' PAGE '
                r3.append(instr)
                p.append(r3)
                r4 = _OE('w:r')
                fld_end = _OE('w:fldChar')
                fld_end.set(_qn('w:fldCharType'), 'end')
                r4.append(fld_end)
                p.append(r4)
            else:
                r = _OE('w:r')
                r.append(rpr)
                t = _OE('w:t')
                t.set(_qn('xml:space'), 'preserve')
                t.text = text_or_fld
                r.append(t)
                p.append(r)
            tc.append(p)
            return tc
        row.append(_footer_cell('Confidential', 'left', 3120))
        row.append(_footer_cell('', 'center', 3120, use_page_field=True))
        row.append(_footer_cell(org_name or '', 'right', 3120))
        tbl.append(row)
        p_std._p.addnext(tbl)

        # Page break after title page
        doc.add_page_break()

        # --- Table of Contents ---
        toc_heading = doc.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_toc_h = toc_heading.add_run('Table of Contents')
        run_toc_h.bold = True
        run_toc_h.font.size = Pt(16)

        # Native Word TOC field - user can right-click > Update Field
        toc_para = doc.add_paragraph()
        toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_toc = toc_para.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        fld_char_begin.set(qn('w:dirty'), 'true')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run_toc._r.append(fld_char_begin)
        run_toc._r.append(instr_text)
        run_toc._r.append(fld_char_end)

        doc.add_page_break()

        # --- Report sections ---
        secs = d.get('sections', [])
        for sec in secs:
            sec_lvl = sec.get('heading_level', 1)
            sec_lvl = sec_lvl if isinstance(sec_lvl, int) and 1 <= sec_lvl <= 9 else 1
            doc.add_heading(sec.get('title', ''), sec_lvl)
            stitle = sec.get('title', '').lower()
            if has_poc and ('appendix b' in stitle or 'points of contact' in stitle):
                doc.add_paragraph('The primary point of contact for this assessment was:')
                poc_name = (poc_first + ' ' + poc_last).strip()
                p_name = doc.add_paragraph()
                p_name.add_run(poc_name).bold = True
                if poc_email:
                    doc.add_paragraph(poc_email)
                if poc_phone:
                    doc.add_paragraph(poc_phone)
            content = sec.get('content', '')
            if content:
                _add_html_to_docx(doc, content, base_heading=sec_lvl + 1)
        if findings:
            doc.add_heading('Findings', 2)
            for f in findings:
                title = f.get('title', '')
                sev = f.get('severity', '')
                doc.add_heading(title + ' [' + sev + ']', 3)
                for lbl, key in [
                    ('Observation', 'description'),
                    ('Discussion', 'discussion'),
                    ('Recommendations', 'recommendation'),
                    ('References', 'refs'),
                ]:
                    html_val = f.get(key, '')
                    if not html_val:
                        continue
                    p2 = doc.add_paragraph()
                    p2.add_run(lbl + ':').bold = True
                    _add_html_to_docx(doc, html_val, base_heading=4)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        data = buf.read()
    except Exception as e:
        raise HTTPException(500, 'DOCX generation failed: ' + str(e))
    fname = (r.title or 'report').replace(' ', '_') + '.docx'
    ct = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return Response(content=data, media_type=ct,
        headers={'Content-Disposition': 'attachment; filename="' + fname + '"'})
